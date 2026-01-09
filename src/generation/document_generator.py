"""
Document Generator - Clean Tag-Based Implementation

Replaces {{ tags }} in the template with corresponding values.
All formatting (bold, highlighting) is handled within the template itself.
"""

import re
from pathlib import Path
from dataclasses import dataclass, field
from typing import List, Dict, Any, Optional
from docx import Document
from docx.table import Table


# --- Data Models ---

@dataclass
class RentRow:
    """Represents a single row in the rent schedule."""
    lease_year_start: int
    lease_year_end: int
    per_sqft: float
    per_annum: float
    per_month: float


@dataclass
class LeaseGenerationInput:
    """Input data for lease document generation."""
    tenant_name: str = ""
    tenant_address: str = ""
    indemnifier_name: str = ""
    indemnifier_address: str = ""
    premises_unit: str = ""
    rentable_area: str = ""
    lease_date: str = ""
    initial_term: str = ""
    renewal_option_count: int = 0
    renewal_option_years: int = 0
    possession_date: str = ""
    fixturing_period: str = ""
    offer_to_lease_date: str = ""
    indemnity_date: str = ""
    rent_schedule: List[RentRow] = field(default_factory=list)
    deposit: str = ""
    tenant_improvement_allowance: str = ""
    permitted_use: str = ""
    trade_name: str = ""
    exclusive_use: str = ""
    radius_restriction: str = ""


# --- Helper Functions ---

def number_to_word(n: int) -> str:
    """Convert an integer to its English word representation (0-200)."""
    ones = ["zero", "one", "two", "three", "four", "five", "six", "seven", 
            "eight", "nine", "ten", "eleven", "twelve", "thirteen", "fourteen",
            "fifteen", "sixteen", "seventeen", "eighteen", "nineteen"]
    tens = ["", "", "twenty", "thirty", "forty", "fifty", "sixty", "seventy", 
            "eighty", "ninety"]
    
    if n < 0 or n > 200:
        return str(n)
    if n < 20:
        return ones[n]
    if n < 100:
        if n % 10 == 0:
            return tens[n // 10]
        return f"{tens[n // 10]}-{ones[n % 10]}"
    if n == 100:
        return "one hundred"
    if n == 200:
        return "two hundred"
    if n <= 119:
        return f"one hundred {ones[n - 100]}"
    if n < 200:
        remainder = n - 100
        if remainder % 10 == 0:
            return f"one hundred {tens[remainder // 10]}"
        return f"one hundred {tens[remainder // 10]}-{ones[remainder % 10]}"
    return str(n)


def parse_date(date_str: str) -> Dict[str, str]:
    """Parse a date string into components."""
    result = {
        "day": "", "day_ordinal": "", "month": "", 
        "month_num": "", "year": "", "year_short": ""
    }
    
    if not date_str:
        return result
    
    months = {
        "january": ("January", "01"), "february": ("February", "02"),
        "march": ("March", "03"), "april": ("April", "04"),
        "may": ("May", "05"), "june": ("June", "06"),
        "july": ("July", "07"), "august": ("August", "08"),
        "september": ("September", "09"), "october": ("October", "10"),
        "november": ("November", "11"), "december": ("December", "12")
    }
    
    ordinals = {1: "1st", 2: "2nd", 3: "3rd", 21: "21st", 22: "22nd", 23: "23rd", 31: "31st"}
    
    # Try "Month Day, Year" format
    match = re.match(r'(\w+)\s+(\d{1,2}),?\s*(\d{4})', date_str)
    if match:
        month_name, day, year = match.groups()
        month_key = month_name.lower()
        if month_key in months:
            day_int = int(day)
            result["day"] = day
            result["day_ordinal"] = ordinals.get(day_int, f"{day}th")
            result["month"], result["month_num"] = months[month_key]
            result["year"] = year
            result["year_short"] = year[-2:]
            return result
    
    # Try "Day Month Year" format
    match = re.match(r'(\d{1,2})\s+(\w+)\s+(\d{4})', date_str)
    if match:
        day, month_name, year = match.groups()
        month_key = month_name.lower()
        if month_key in months:
            day_int = int(day)
            result["day"] = day
            result["day_ordinal"] = ordinals.get(day_int, f"{day}th")
            result["month"], result["month_num"] = months[month_key]
            result["year"] = year
            result["year_short"] = year[-2:]
    
    return result


def extract_number(text: str) -> Optional[int]:
    """Extract first integer from a string."""
    match = re.search(r'(\d+)', str(text))
    return int(match.group(1)) if match else None


# --- Document Generator ---

class DocumentGenerator:
    """Generates lease documents by replacing {{ tags }} with values."""
    
    DEFAULT_TEMPLATE_PATH = Path("data/raw_pdfs/Lease Precedent 11939 240 Street MR Jun 19, 2020 -- with tags.docx")
    
    def __init__(self, template_path: Path = None):
        self.template_path = template_path or self.DEFAULT_TEMPLATE_PATH
        if not self.template_path.exists():
            raise FileNotFoundError(f"Template not found: {self.template_path}")
    
    def generate(self, input_data: LeaseGenerationInput, output_path: Path) -> Path:
        """Generate a populated lease document."""
        print(f"[DEBUG] Using template: {self.template_path}")
        doc = Document(str(self.template_path))
        
        # Build tag -> value mapping
        tags = self._build_tag_mapping(input_data)
        print(f"[DEBUG] Generated {len(tags)} tags")
        print(f"[DEBUG] Sample tags: tenant_name='{tags.get('{{ tenant_name }}', 'N/A')}'")
        
        # Replace in all paragraphs (body, headers, footers)
        self._replace_tags_in_document(doc, tags)
        
        # Handle rent schedule table separately (dynamic rows)
        self._populate_rent_schedule(doc, input_data.rent_schedule)
        
        # Save
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        print(f"[DEBUG] Saved to: {output_path}")
        return output_path
    
    def _build_tag_mapping(self, data: LeaseGenerationInput) -> Dict[str, str]:
        """Build a dictionary mapping {{ tag }} -> replacement value."""
        tags = {}
        
        # Basic fields
        tags["{{ tenant_name }}"] = data.tenant_name
        tags["{{ tenant_address }}"] = data.tenant_address
        tags["{{ indemnifier_name }}"] = data.indemnifier_name
        tags["{{ indemnifier_address }}"] = data.indemnifier_address
        # Alias for template typo (idemnifier instead of indemnifier)
        tags["{{ idemnifier_name }}"] = data.indemnifier_name
        tags["{{ idemnifier_address }}"] = data.indemnifier_address
        tags["{{ premises_unit }}"] = data.premises_unit
        tags["{{ rentable_area }}"] = data.rentable_area
        tags["{{ trade_name }}"] = data.trade_name
        tags["{{ permitted_use }}"] = data.permitted_use
        tags["{{ exclusive_use }}"] = data.exclusive_use
        tags["{{ deposit_amount }}"] = data.deposit
        tags["{{ ti_allowance }}"] = data.tenant_improvement_allowance
        
        # Simple dates
        tags["{{ offer_date }}"] = data.offer_to_lease_date
        tags["{{ possession_date }}"] = data.possession_date
        
        # Parsed dates (lease date)
        lease_parts = parse_date(data.lease_date)
        tags["{{ lease_full }}"] = data.lease_date
        tags["{{ lease_day_ordinal }}"] = lease_parts["day_ordinal"]
        tags["{{ lease_month }}"] = lease_parts["month"]
        tags["{{ lease_year_full }}"] = lease_parts["year"]
        tags["{{ lease_year_short }}"] = lease_parts["year_short"]
        
        # Parsed dates (indemnity date)
        indem_parts = parse_date(data.indemnity_date)
        tags["{{ indemnity_full }}"] = data.indemnity_date
        tags["{{ indemnity_day_ordinal }}"] = indem_parts["day_ordinal"]
        tags["{{ indemnity_month }}"] = indem_parts["month"]
        tags["{{ indemnity_year_full }}"] = indem_parts["year"]
        tags["{{ indemnity_year_short }}"] = indem_parts["year_short"]
        
        # Term (word + number pairs)
        term_num = extract_number(data.initial_term)
        if term_num:
            tags["{{ term_years_word }}"] = number_to_word(term_num)
            tags["{{ term_years_num }}"] = str(term_num)
        
        # Fixturing period
        fix_num = extract_number(data.fixturing_period)
        if fix_num:
            tags["{{ fixturing_days_word }}"] = number_to_word(fix_num)
            tags["{{ fixturing_days_num }}"] = str(fix_num)
        
        # Renewal options
        if data.renewal_option_count:
            tags["{{ renewal_count_word }}"] = number_to_word(data.renewal_option_count)
            tags["{{ renewal_count_num }}"] = str(data.renewal_option_count)
        if data.renewal_option_years:
            tags["{{ renewal_years_word }}"] = number_to_word(data.renewal_option_years)
            tags["{{ renewal_years_num }}"] = str(data.renewal_option_years)
        
        # Radius restriction
        rad_num = extract_number(data.radius_restriction)
        if rad_num:
            tags["{{ radius_word }}"] = number_to_word(rad_num)
            tags["{{ radius_num }}"] = str(rad_num)
        
        return tags
    
    def _replace_tags_in_document(self, doc: Document, tags: Dict[str, str]):
        """Replace all {{ tags }} throughout the document."""
        # Body paragraphs
        for para in doc.paragraphs:
            self._replace_tags_in_paragraph(para, tags)
        
        # Tables (non-rent-schedule cells)
        for table in doc.tables:
            if not self._is_rent_schedule_table(table):
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            self._replace_tags_in_paragraph(para, tags)
        
        # Headers and footers
        for section in doc.sections:
            for header in [section.header, section.first_page_header, section.even_page_header]:
                if header:
                    for para in header.paragraphs:
                        self._replace_tags_in_paragraph(para, tags)
            for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                if footer:
                    for para in footer.paragraphs:
                        self._replace_tags_in_paragraph(para, tags)
    
    def _replace_tags_in_paragraph(self, para, tags: Dict[str, str]):
        """Replace tags in a single paragraph, preserving run formatting."""
        if "{{" not in para.text:
            return
        
        for tag, value in tags.items():
            if not value or tag not in para.text:
                continue
            
            # Try to replace within individual runs first (preserves formatting perfectly)
            replaced = False
            for run in para.runs:
                if tag in run.text:
                    run.text = run.text.replace(tag, value)
                    replaced = True
            
            # If tag spans multiple runs, we need to find and reconstruct
            if not replaced and tag in para.text:
                # Build a map of character positions to runs
                runs_text = [(run, run.text) for run in para.runs]
                full_text = "".join(t for _, t in runs_text)
                tag_start = full_text.find(tag)
                
                if tag_start >= 0:
                    tag_end = tag_start + len(tag)
                    # Find which runs contain the tag
                    pos = 0
                    for run, text in runs_text:
                        run_start = pos
                        run_end = pos + len(text)
                        
                        # If this run contains start of tag
                        if run_start <= tag_start < run_end:
                            # Replace from tag_start to end of this run with value + remaining
                            before = text[:tag_start - run_start]
                            after_in_run = text[min(tag_end - run_start, len(text)):]
                            run.text = before + value + after_in_run
                            # Clear subsequent runs that were part of the tag
                            clear_remaining = tag_end - run_end
                            if clear_remaining > 0:
                                for next_run, next_text in runs_text[runs_text.index((run, text)) + 1:]:
                                    if clear_remaining >= len(next_text):
                                        next_run.text = ""
                                        clear_remaining -= len(next_text)
                                    else:
                                        next_run.text = next_text[clear_remaining:]
                                        break
                            break
                        pos = run_end
    
    def _is_rent_schedule_table(self, table: Table) -> bool:
        """Check if table is the rent schedule based on header keywords."""
        if not table.rows:
            return False
        first_row_text = " ".join(cell.text.lower() for cell in table.rows[0].cells)
        keywords = ["lease year", "per square foot", "per annum", "per month"]
        return sum(1 for kw in keywords if kw in first_row_text) >= 2
    
    def _populate_rent_schedule(self, doc: Document, rent_schedule: List[RentRow]):
        """Populate the rent schedule table with data rows."""
        if not rent_schedule:
            return
        
        for table in doc.tables:
            if not self._is_rent_schedule_table(table):
                continue
            
            # Skip header row (index 0), populate data rows
            for i, rent_row in enumerate(rent_schedule):
                row_idx = 1 + i
                
                # Add row if needed
                if row_idx >= len(table.rows):
                    table.add_row()
                
                row = table.rows[row_idx]
                if len(row.cells) >= 4:
                    year_range = f"{rent_row.lease_year_start}-{rent_row.lease_year_end}" \
                        if rent_row.lease_year_start != rent_row.lease_year_end \
                        else str(rent_row.lease_year_start)
                    row.cells[0].text = year_range
                    row.cells[1].text = f"${rent_row.per_sqft:.2f}"
                    row.cells[2].text = f"${rent_row.per_annum:,.2f}"
                    row.cells[3].text = f"${rent_row.per_month:,.2f}"
            break


# --- Factory Function ---

def generate_lease_document(
    input_data: Dict[str, Any],
    output_filename: str = "generated_lease.docx",
    output_dir: str = "output"
) -> Path:
    """Generate a lease document from a dictionary of input data."""
    rent_rows = [
        RentRow(
            lease_year_start=row.get("lease_year_start", 1),
            lease_year_end=row.get("lease_year_end", 1),
            per_sqft=row.get("per_sqft", 0.0),
            per_annum=row.get("per_annum", 0.0),
            per_month=row.get("per_month", 0.0),
        )
        for row in input_data.get("rent_schedule", [])
    ]
    
    lease_input = LeaseGenerationInput(
        tenant_name=input_data.get("tenant_name", ""),
        tenant_address=input_data.get("tenant_address", ""),
        indemnifier_name=input_data.get("indemnifier_name", ""),
        indemnifier_address=input_data.get("indemnifier_address", ""),
        premises_unit=input_data.get("premises_unit", ""),
        rentable_area=input_data.get("rentable_area", ""),
        lease_date=input_data.get("lease_date", ""),
        initial_term=input_data.get("initial_term", ""),
        renewal_option_count=input_data.get("renewal_option_count", 0),
        renewal_option_years=input_data.get("renewal_option_years", 0),
        possession_date=input_data.get("possession_date", ""),
        fixturing_period=input_data.get("fixturing_period", ""),
        offer_to_lease_date=input_data.get("offer_to_lease_date", ""),
        indemnity_date=input_data.get("indemnity_date", ""),
        rent_schedule=rent_rows,
        deposit=input_data.get("deposit", ""),
        tenant_improvement_allowance=input_data.get("tenant_improvement_allowance", ""),
        permitted_use=input_data.get("permitted_use", ""),
        trade_name=input_data.get("trade_name", ""),
        exclusive_use=input_data.get("exclusive_use", ""),
        radius_restriction=input_data.get("radius_restriction", ""),
    )
    
    generator = DocumentGenerator()
    output_path = Path(output_dir) / output_filename
    return generator.generate(lease_input, output_path)


# --- Test Block ---

if __name__ == "__main__":
    sample_input = {
        "tenant_name": "ACME Corporation Ltd.",
        "tenant_address": "123 Main Street, Vancouver, BC V6B 1A1",
        "indemnifier_name": "John Smith",
        "indemnifier_address": "456 Oak Avenue, Vancouver, BC V6C 2B2",
        "premises_unit": "Unit #200",
        "rentable_area": "2,500 square feet",
        "lease_date": "January 15, 2026",
        "initial_term": "five (5) years",
        "renewal_option_count": 2,
        "renewal_option_years": 5,
        "possession_date": "February 1, 2026",
        "fixturing_period": "60 days",
        "offer_to_lease_date": "December 1, 2025",
        "indemnity_date": "January 15, 2026",
        "rent_schedule": [
            {"lease_year_start": 1, "lease_year_end": 2, "per_sqft": 35.00, "per_annum": 87500.00, "per_month": 7291.67},
            {"lease_year_start": 3, "lease_year_end": 5, "per_sqft": 38.00, "per_annum": 95000.00, "per_month": 7916.67},
        ],
        "deposit": "$15,000.00",
        "tenant_improvement_allowance": "$25.00 per square foot",
        "permitted_use": "General office use",
        "trade_name": "ACME Solutions",
        "exclusive_use": "None",
        "radius_restriction": "5 km",
    }
    
    try:
        output = generate_lease_document(sample_input, "test_clean.docx")
        print(f"✅ Generated: {output}")
    except FileNotFoundError as e:
        print(f"⚠️ Template not found: {e}")
    except Exception as e:
        print(f"❌ Error: {e}")
